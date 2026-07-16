import re
import shutil
import subprocess
import tempfile
import zipfile
from collections.abc import Callable
from pathlib import Path
from typing import Any
from uuid import uuid4
from xml.etree import ElementTree

from docx import Document
from fastapi import status
from pydantic import ValidationError

from app.core.config import Settings, get_settings
from app.core.errors import AppError, ErrorCode
from app.core.http_status import HTTP_422_UNPROCESSABLE_CONTENT
from app.schemas.resume import StructuredResumeData
from app.services.llm import LLMClient

MAX_RESUME_BYTES = 10 * 1024 * 1024
MAX_DOCX_MEMBERS = 2_048
MAX_DOCX_MEMBER_BYTES = 16 * 1024 * 1024
MAX_DOCX_UNCOMPRESSED_BYTES = 64 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 200
SUPPORTED_EXTENSIONS = {".doc", ".docx"}

ProjectConverter = Callable[[Path, Path], Path]


class ResumeParserService:
    def __init__(
        self,
        llm_client: LLMClient,
        settings: Settings | None = None,
        converter: ProjectConverter | None = None,
    ) -> None:
        self.llm_client = llm_client
        self.settings = settings or get_settings()
        self.converter = converter or convert_doc_to_docx

    def parse(self, original_path: Path) -> dict[str, Any]:
        if original_path.suffix.lower() == ".doc":
            with tempfile.TemporaryDirectory(
                dir=original_path.parent,
                prefix=f"{original_path.stem}-converted-",
            ) as temp_dir:
                parse_path = self.converter(original_path, Path(temp_dir))
                _validate_converted_path(parse_path, Path(temp_dir))
                resume_text = extract_docx_text(parse_path)
        else:
            resume_text = extract_docx_text(original_path)

        if not resume_text.strip():
            raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)
        _validate_resume_text_length(resume_text, self.settings.resume_max_text_chars)

        project_candidates = extract_project_title_candidates(resume_text)
        parse_text = _resume_text_with_project_constraints(resume_text, project_candidates)
        structured_data = self.llm_client.parse_resume(parse_text)
        try:
            validated = StructuredResumeData.model_validate(structured_data).model_dump()
            return ensure_project_completeness(resume_text, validated)
        except ValidationError as exc:
            raise AppError(
                ErrorCode.RESUME_PARSE_FAILED,
                HTTP_422_UNPROCESSABLE_CONTENT,
            ) from exc


def project_root() -> Path:
    return Path(__file__).resolve().parents[3]


def resolve_upload_dir(settings: Settings | None = None) -> Path:
    configured = Path((settings or get_settings()).upload_dir)
    if not configured.is_absolute():
        configured = project_root() / configured
    return configured


def validate_resume_extension(filename: str) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise AppError(ErrorCode.INVALID_UPLOAD_TYPE, status.HTTP_400_BAD_REQUEST)
    return extension


def store_resume_upload(filename: str, upload_dir: Path, content: bytes) -> Path:
    extension = validate_resume_extension(filename)
    upload_dir.mkdir(parents=True, exist_ok=True)
    for _attempt in range(3):
        candidate = upload_dir / f"{uuid4().hex}{extension}"
        try:
            with candidate.open("xb") as target:
                target.write(content)
            return candidate
        except FileExistsError:
            continue
    raise RuntimeError("unable to allocate a unique resume upload path")


def extract_docx_text(path: Path) -> str:
    validate_docx_archive(path)
    try:
        document = Document(str(path))
    except Exception as exc:
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT) from exc

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append("\t".join(cells))
    text = "\n".join(paragraphs)
    if text.strip():
        return text
    return extract_docx_xml_text(path)


def extract_docx_xml_text(path: Path) -> str:
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    xml_parts = (
        "word/document.xml",
        "word/header1.xml",
        "word/header2.xml",
        "word/header3.xml",
        "word/footer1.xml",
        "word/footer2.xml",
        "word/footer3.xml",
    )
    paragraphs: list[str] = []

    try:
        with zipfile.ZipFile(path) as archive:
            validate_docx_archive(path, archive=archive)
            for part_name in xml_parts:
                if part_name not in archive.namelist():
                    continue
                root = ElementTree.fromstring(_read_archive_member(archive, part_name))
                for paragraph in root.findall(".//w:p", namespaces):
                    if paragraph.findall(".//w:p", namespaces):
                        continue
                    text = "".join(
                        node.text or ""
                        for node in paragraph.findall(".//w:t", namespaces)
                    ).strip()
                    if text:
                        paragraphs.append(text)
    except (ElementTree.ParseError, OSError, zipfile.BadZipFile) as exc:
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT) from exc

    return "\n".join(paragraphs)


def validate_docx_archive(path: Path, *, archive: zipfile.ZipFile | None = None) -> None:
    try:
        if archive is None:
            with zipfile.ZipFile(path) as opened_archive:
                _validate_docx_members(opened_archive.infolist())
        else:
            _validate_docx_members(archive.infolist())
    except (OSError, zipfile.BadZipFile) as exc:
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT) from exc


def _validate_docx_members(members: list[zipfile.ZipInfo]) -> None:
    if not members or len(members) > MAX_DOCX_MEMBERS:
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)

    total_size = 0
    for member in members:
        member_path = Path(member.filename.replace("\\", "/"))
        if member.flag_bits & 0x1 or member_path.is_absolute() or ".." in member_path.parts:
            raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)
        if member.is_dir():
            continue

        total_size += member.file_size
        if member.file_size > MAX_DOCX_MEMBER_BYTES or total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
            raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)
        if member.file_size > 0:
            compressed_size = max(1, member.compress_size)
            if member.file_size / compressed_size > MAX_DOCX_COMPRESSION_RATIO:
                raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)


def _read_archive_member(archive: zipfile.ZipFile, member_name: str) -> bytes:
    with archive.open(member_name) as source:
        content = source.read(MAX_DOCX_MEMBER_BYTES + 1)
    if len(content) > MAX_DOCX_MEMBER_BYTES:
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)
    return content


DATE_RANGE_HEADING = re.compile(
    r"^\s*(?:19|20)\d{2}(?:[./年-]\d{1,2})?\s*"
    r"(?:"
    r"[-–—~～]\s*(?:(?:19|20)\d{2}(?:[./年-]\d{1,2})?|至今|现在|present)"
    r"|至今|至现在|到现在"
    r")"
    r"\s*(?P<title>.+?)\s*$",
    re.IGNORECASE,
)
PROJECT_TITLE_TRAILING_MARKERS = re.compile(
    r"\s+(?:github|gitee|gitlab|项目链接|源码链接|demo)\s*$",
    re.IGNORECASE,
)
NON_PROJECT_HEADING_MARKERS = (
    "大学",
    "学院",
    "本科",
    "硕士",
    "博士",
    "学士",
    "有限公司",
    "实习生",
    "工程师",
    "开发岗",
    "产品经理",
)


def extract_project_title_candidates(resume_text: str) -> list[str]:
    candidates: list[str] = []
    for raw_line in resume_text.splitlines():
        line = " ".join(raw_line.split())
        matched = DATE_RANGE_HEADING.match(line)
        if matched is None:
            continue
        title = PROJECT_TITLE_TRAILING_MARKERS.sub("", matched.group("title")).strip(" -—|：:")
        if not title or len(title) > 160:
            continue
        normalized = _normalize_resume_text(title)
        if normalized and all(_normalize_resume_text(item) != normalized for item in candidates):
            candidates.append(title)
    return candidates


def ensure_project_completeness(
    resume_text: str,
    structured_data: dict[str, Any],
) -> dict[str, Any]:
    projects = structured_data.get("project_experience")
    if not isinstance(projects, list):
        projects = []
        structured_data["project_experience"] = projects

    excluded_headings = _non_project_titles(structured_data)
    existing_titles = [_project_title(project) for project in projects]
    for candidate in extract_project_title_candidates(resume_text):
        if _is_non_project_heading(candidate, excluded_headings):
            continue
        if any(_same_resume_title(candidate, title) for title in existing_titles if title):
            continue
        projects.append({"name": candidate})
        existing_titles.append(candidate)
    return structured_data


def _resume_text_with_project_constraints(resume_text: str, candidates: list[str]) -> str:
    likely_projects = [
        title
        for title in candidates
        if not any(marker in title for marker in NON_PROJECT_HEADING_MARKERS)
    ]
    if not likely_projects:
        return resume_text
    candidate_lines = "\n".join(f"- {title}" for title in likely_projects)
    return (
        f"{resume_text}\n\n"
        "【项目完整性约束】以下是从原文日期标题中识别到的项目候选。"
        "请逐项核对，属于项目的条目必须全部保留在 project_experience，"
        "即使原文暂时只有项目名称也不能省略：\n"
        f"{candidate_lines}"
    )


def _non_project_titles(structured_data: dict[str, Any]) -> list[str]:
    titles: list[str] = []
    for section, keys in (
        ("education", ("school", "institution", "name")),
        ("work_experience", ("company", "organization", "name")),
    ):
        entries = structured_data.get(section)
        if not isinstance(entries, list):
            continue
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            titles.extend(
                str(entry.get(key) or "").strip()
                for key in keys
                if str(entry.get(key) or "").strip()
            )
    return titles


def _is_non_project_heading(candidate: str, excluded_headings: list[str]) -> bool:
    normalized = _normalize_resume_text(candidate)
    if any(marker in candidate for marker in NON_PROJECT_HEADING_MARKERS):
        return True
    return any(
        excluded
        and (excluded in normalized or normalized in excluded)
        for excluded in (_normalize_resume_text(item) for item in excluded_headings)
    )


def _project_title(project: Any) -> str:
    if not isinstance(project, dict):
        return ""
    for key in ("name", "project_name", "title"):
        value = str(project.get(key) or "").strip()
        if value:
            return value
    return ""


def _same_resume_title(left: str, right: str) -> bool:
    normalized_left = _normalize_resume_text(left)
    normalized_right = _normalize_resume_text(right)
    return bool(
        normalized_left
        and normalized_right
        and (
            normalized_left == normalized_right
            or normalized_left in normalized_right
            or normalized_right in normalized_left
        )
    )


def _normalize_resume_text(value: str) -> str:
    return "".join(char for char in value.casefold() if char.isalnum())


def _validate_resume_text_length(resume_text: str, max_chars: int) -> None:
    if len(resume_text) <= max(1, max_chars):
        return
    raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)


def _validate_converted_path(path: Path, output_dir: Path) -> None:
    try:
        converted_path = path.resolve(strict=True)
        converted_root = output_dir.resolve(strict=True)
    except OSError as exc:
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT) from exc
    if (
        converted_path.suffix.lower() != ".docx"
        or not converted_path.is_relative_to(converted_root)
    ):
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)


def convert_doc_to_docx(input_path: Path, output_dir: Path) -> Path:
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if libreoffice is None:
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{input_path.stem}.docx"
    if output_path.exists():
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)
    try:
        subprocess.run(
            [
                libreoffice,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(output_dir),
                str(input_path),
            ],
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            timeout=max(1, get_settings().resume_conversion_timeout_seconds),
        )
    except (OSError, subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT) from exc

    if not output_path.exists():
        raise AppError(ErrorCode.RESUME_PARSE_FAILED, HTTP_422_UNPROCESSABLE_CONTENT)
    return output_path
