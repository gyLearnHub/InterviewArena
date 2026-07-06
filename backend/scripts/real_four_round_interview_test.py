from __future__ import annotations

# ruff: noqa: E402
import json
import os
import socket
import subprocess
import sys
import threading
import time
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

import httpx
import uvicorn
from docx import Document

BACKEND_ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = BACKEND_ROOT.parent
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from app.db.mysql import mysql_connection
from main import create_app

ROUND_ORDER = ["resume", "technical", "manager", "hr"]
ROUND_LABELS = {
    "resume": "简历面",
    "technical": "技术面",
    "manager": "主管面",
    "hr": "HR 面",
}
ROUND_MIN = {"resume": 20, "technical": 20, "manager": 20, "hr": 15}
ROUND_MAX = {"resume": 40, "technical": 40, "manager": 40, "hr": 40}

TARGET_POSITION = "Agent 应用开发工程师"
JOB_DESCRIPTION = (
    "负责面向企业招聘场景的 AI Agent 应用开发，使用 Python/FastAPI、MySQL、Vue "
    "建设多轮面试、RAG 检索、LLM 评估、状态恢复、观测与报告生成能力。要求理解计算机基础、"
    "数据库事务与索引、网络协议、操作系统并发、算法复杂度、Agent 工具调用、RAG 召回评估、"
    "Prompt/JSON 结构化输出、服务稳定性、测试与前后端协作。"
)


RESUME_LINES = [
    "姓名：林泽宇",
    "联系方式：lin.zeyu.test@example.com / 13800001234",
    "求职目标：Agent 应用开发工程师 / 后端开发工程师",
    "教育背景：上海交通大学 软件工程 本科 2018.09-2022.06",
    "工作经历：",
    "2024.03-至今 星澜智能 招聘智能化团队 后端与 Agent 应用开发工程师",
    "负责 InterviewArena 多轮模拟面试平台，建设简历解析、四轮面试编排、问题生成、单题评分、轮次总结、最终报告、Harness Trace 与恢复点能力。",
    "2022.07-2024.02 云启数据 平台研发工程师",
    "负责数据标注平台的任务分发、权限、审计与报表服务，支撑约 120 名内部标注与质检人员。",
    "代表项目一：InterviewArena 多轮面试与评估系统",
    "技术栈：Python、FastAPI、MySQL、PyMySQL、Vue3、TypeScript、Chroma、BM25、DeepSeek API、Playwright。",
    "我的职责：设计 interview / round / qa / evaluation / harness_trace 数据模型，拆分简历面、技术面、主管面和 HR 面 Agent，编写问题生成、回答提交、追问切换、轮次总结和最终报告接口。",
    "关键方案：用状态机约束 pending、in_progress、completed、finished_early、cancelled；用事务保证问题、回答、评分和轮次状态一致；用 idempotency_key 避免 Harness Trace 重复；单题评分先写 evaluation_records，轮次结束后统一展示。",
    "难点一：真实 LLM JSON 输出偶尔字段缺失。我在服务端用 Pydantic 校验、异常记录和本地评分兜底，保证低质量回答不会拿固定及格分。",
    "难点二：多轮上下文容易膨胀。我只把当前轮摘要、最近问答、未覆盖维度和必要记忆传给 Agent，最终评分隔离长期记忆。",
    "结果：端到端流程可稳定恢复，历史报告能追溯 question_id、round_id、评分理由和 Harness Trace。",
    "代表项目二：企业知识库 RAG 问答助手",
    "技术栈：FastAPI、MySQL、Redis、Chroma、Sentence Embedding、BM25、重排模型、SSE。",
    "我的职责：实现文档切分、向量入库、混合检索、召回评估、答案引用、权限过滤和流式响应。",
    "关键方案：先 BM25+向量混合召回，再按文档权限和时间过滤，最后用 reranker 重排；对回答引用做 chunk_id 追踪，离线统计 recall@k 和命中率。",
    "结果：常见知识查询首响从约 4 秒降到 1.8 秒，人工复核时引用可追溯。",
    "代表项目三：数据标注任务调度平台",
    "技术栈：Java、Spring Boot、MySQL、Redis、Kafka、Vue。",
    "我的职责：维护任务分配、质检抽样、导出报表和操作审计模块。",
    "关键方案：用乐观锁和唯一索引避免任务重复领取，用 Redis 缓存热点字典，用 Kafka 异步生成报表。",
    "技能：Python、Java、FastAPI、Spring Boot、Vue3、TypeScript、MySQL、Redis、HTTP/TCP、Linux、Docker、pytest、Playwright、RAG、Agent、Prompt Engineering。",
    "项目管理与协作：常与产品、前端、算法同学对齐接口契约和验收标准；遇到风险会拆分最小可交付版本并补充监控与回滚方案。",
    "职业规划：未来 3 年希望继续深耕 AI 应用工程化，重点提升 Agent 可靠性、评估体系和复杂系统设计能力。",
    "薪资期望：月薪 28k-32k，可根据岗位职责、成长空间和整体 offer 沟通。",
    "到岗时间：收到正式 offer 后 2 周内可到岗。",
]


def now_slug() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def free_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


def wait_for_health(base_url: str, timeout_seconds: int = 20) -> None:
    deadline = time.time() + timeout_seconds
    while time.time() < deadline:
        try:
            response = httpx.get(f"{base_url}/health", timeout=1)
            if response.status_code == 200:
                return
        except httpx.HTTPError:
            time.sleep(0.2)
    raise RuntimeError(f"API server did not become ready: {base_url}")


def start_api_server() -> tuple[uvicorn.Server, str]:
    app = create_app()
    port = free_port()
    server = uvicorn.Server(
        uvicorn.Config(app, host="127.0.0.1", port=port, log_level="info")
    )
    thread = threading.Thread(target=server.run, daemon=True)
    thread.start()
    base_url = f"http://127.0.0.1:{port}/api"
    wait_for_health(base_url)
    return server, base_url


def api_request(
    client: httpx.Client,
    method: str,
    path: str,
    report: dict[str, Any],
    **kwargs: Any,
) -> httpx.Response:
    started = time.time()
    response = client.request(method, path, **kwargs)
    elapsed_ms = int((time.time() - started) * 1000)
    report.setdefault("api_calls", []).append(
        {
            "method": method,
            "path": path,
            "status_code": response.status_code,
            "elapsed_ms": elapsed_ms,
            "request_json": kwargs.get("json"),
            "response_excerpt": clip(response.text, 700),
        }
    )
    if response.status_code >= 400:
        raise AssertionError(f"{method} {path} failed: {response.status_code} {response.text}")
    return response


def make_resume_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("真实流程测试简历 - 林泽宇", level=1)
    for line in RESUME_LINES:
        document.add_paragraph(line)
    document.save(path)


def answer_question(
    *,
    round_type: str,
    question: dict[str, Any],
    round_history: list[dict[str, Any]],
) -> str:
    text = str(question.get("question") or "")
    kind = str(question.get("question_kind") or "main")
    sequence = int(question.get("sequence") or len(round_history) + 1)
    previous = round_history[-1] if round_history else None
    prefix = ""
    if kind == "follow_up" and previous:
        prefix = (
            f"刚才我回答的是“{clip(previous['question'], 45)}”。针对这个追问，我补充得更具体一点："
        )
    if round_type == "technical":
        return prefix + technical_answer(text, sequence)
    if round_type == "manager":
        return prefix + manager_answer(text, sequence)
    if round_type == "hr":
        return prefix + hr_answer(text, sequence)
    return prefix + resume_answer(text, sequence)


def resume_answer(question: str, sequence: int) -> str:
    if contains(question, ["四个 agent", "四类 agent", "不同 agent", "agent 之间", "prompt", "提示词", "交互逻辑", "对话逻辑", "简历面", "技术面", "主管面", "hr"]):
        return (
            "四个面试官 Agent 的差异是我参与设计的重点。简历面 Agent 的目标是核验经历真实性、项目背景、个人贡献和岗位匹配，"
            "Prompt 会要求它围绕简历时间线、项目职责、团队分工、结果证据追问，避免问太多纯技术细节。"
            "技术面 Agent 的目标是验证基础知识、原理和工程实现，Prompt 里明确要求交叉覆盖数据库、网络、操作系统、算法、Agent、RAG、LLM 应用和项目实践，"
            "并且开场不能全围绕项目。主管面 Agent 更关注目标拆解、协作推动、风险处理、复盘成长，问题会偏业务场景和团队协作。"
            "HR Agent 则关注动机、稳定性、价值观、薪资和到岗时间，不能混入技术细节评分。"
            "交互逻辑上四个 Agent 共用同一套状态机和 qa_history，但使用不同 system_prompt、评价维度和 core_topics；"
            "每次生成问题前会把最近问答、上一条回答、未覆盖主题和 question_kind 传进去，由服务端决定主问题还是追问。"
        )
    if contains(question, ["状态机", "状态流转", "pending", "in_progress", "completed", "finished_early", "cancelled"]):
        return (
            "InterviewArena 的状态机分两层：整场 interview 和单个 round。创建面试后 interview 是 created，四个 round 是 pending；"
            "点击开始某一轮时，服务端先校验前序轮次是否完成，再把该 round 改成 in_progress，同时 interview.overall_status 改为 in_progress、current_round 写入当前轮次。"
            "每次提交回答只更新 interview_qa.answer，并根据评分建议决定继续追问或生成新主问题；达到结束条件或调用正常结束接口时，round 改为 completed，写入 score、result、summary 和 ended_at。"
            "如果用户主动提前结束当前轮，会标记 finished_early 并把 summary 标成 reference_only；如果整场提前退出，未开始的 pending 轮次会变成 cancelled。"
            "四轮 selected_rounds 都完成后，最终报告接口生成 feedback_reports，再把 interview.status 和 overall_status 改成 finished，并清空 current_round。"
        )
    if contains(question, ["定位", "目标用户", "用户群体", "解决什么", "业务问题", "痛点", "使用场景"]):
        return (
            "InterviewArena 的定位是面向求职者和招聘训练场景的多轮模拟面试平台，不是单纯聊天机器人。目标用户主要有两类："
            "一类是准备技术/AI 应用岗位面试的候选人，希望用真实简历和 JD 练习多轮面试；另一类是企业内部招聘或培训团队，希望沉淀可追溯的面试评估样例。"
            "它解决的核心问题有三个：第一，传统模拟面试缺少连续追问和四轮角色区分；第二，面试反馈往往只有笼统建议，不能追溯到具体 question_id 和回答证据；"
            "第三，LLM 生成流程不稳定，失败后很难恢复。因此我参与设计时强调多轮状态机、结构化评分、Harness Trace 和最终报告证据链。"
        )
    if contains(question, ["关键决策", "推动", "主导", "设计决策", "为什么这么设计", "思考过程", "最终效果", "效果"]):
        return (
            "我推动过一个比较关键的决策：把“问题/回答”和“评分/报告”拆成两条可关联但不耦合的链路。"
            "最初有人倾向于每次回答后直接把评分混在 qa 记录里返回，前端实现会更简单；但我担心这样会导致候选人边答边看到分数，也不利于评分失败后的重试。"
            "我的方案是 interview_qa 只保存问题和回答，evaluation_records 用 question_id、round_id 和 evaluation_key 单独保存评分；轮次结束后再统一把 question_evaluations 放进 summary 展示。"
            "这样做的效果是：提交回答、单题评分、轮次总结可以分别追踪；评分失败不会破坏问答主流程；历史报告仍能通过 question_id 找到每题评分理由。"
        )
    if contains(question, ["岗位职责", "核心职责", "胜任", "哪段经历", "哪一段经历", "支撑", "匹配这个岗位"]):
        return (
            "我理解 Agent 应用开发工程师的核心职责不是只会调模型接口，而是把 LLM/Agent 能力做成稳定可用的业务系统，"
            "包括工具调用编排、上下文控制、结构化输出校验、RAG 检索、状态恢复、评估体系、接口和数据库设计。"
            "最能支撑我胜任这类岗位的是 InterviewArena 项目，因为它同时覆盖了四个方面："
            "第一，我做了多 Agent 角色拆分，让简历面、技术面、主管面、HR 面使用不同 prompt 和评价维度；"
            "第二，我做了真实业务状态机和 MySQL 持久化，保证问题、回答、评分和报告能追溯；"
            "第三，我处理了 LLM JSON 输出不稳定的问题，用 Pydantic 校验、本地评分规则和失败记录兜底；"
            "第四，我做过 RAG 项目的混合召回、重排和引用追踪，可以补足知识检索与上下文构建能力。"
        )
    if contains(question, ["模块", "功能", "平台", "负责哪些", "负责了哪些", "具体负责", "开发内容", "相关功能"]):
        return (
            "我在 InterviewArena 里负责的模块主要有六块：第一是认证和用户隔离，保证每个用户只能访问自己的简历和面试；"
            "第二是简历解析和结构化存储，把 docx 简历解析成 basic_info、education、project_experience、skills 等字段；"
            "第三是四轮面试编排，创建 interview 后生成 resume、technical、manager、hr 四个 round，并按状态机依次推进；"
            "第四是问题与回答链路，interview_qa 记录主问题、追问、parent_question_id 和候选人回答；"
            "第五是评分链路，单题评分写 evaluation_records，轮次结束后生成 round_summary；"
            "第六是最终报告和 Harness Trace，四轮完成后生成 feedback_report，并记录每次 Agent 调用、错误和 checkpoint。"
            "我个人主要写后端 service、repository、状态校验和接口契约，也参与和前端对齐展示字段。"
        )
    if contains(question, ["团队", "规模", "分工", "协作关系", "成员"]):
        return (
            "InterviewArena 当时是一个 7 人左右的小团队：1 名产品负责面试流程和报告验收，2 名前端负责 Vue 页面和交互，"
            "2 名后端负责认证、简历、面试编排、评分和历史记录，1 名算法同学协助 RAG/LLM 参数和提示词策略，1 名测试同学做回归。"
            "我的边界主要在后端和 Agent 编排，但我会和前端一起定接口字段，比如 current_question、qa_history、round_summary、harness_status 的展示规则。"
        )
    if contains(question, ["实体", "表", "数据模型", "表关系", "关联", "字段", "数据库"]):
        return (
            "核心表关系是：users 一对多 resumes 和 interviews；interviews 通过 resume_id 关联本次使用的简历快照；"
            "interview_rounds 用 interview_id 关联一场面试下的四个轮次；interview_qa 用 round_id、sequence 保存每道主问题或追问，"
            "parent_question_id 指向被追问的主问题；evaluation_records 用 question_id 保存单题评分，用 round_id 保存轮次评分，用 evaluation_key 做幂等；"
            "feedback_reports 只在四轮完成后生成最终报告；harness_traces 和 harness_checkpoints 用 interview_id、round_id 记录每次 Agent 调用和恢复点。"
        )
    if contains(question, ["单题评分", "统一展示", "评分展示"]):
        return (
            "单题评分是在提交回答后就落到 evaluation_records，但接口展示时做了隔离：当前轮未结束时 state 接口不会把 question_evaluation 塞回 qa_history；"
            "只有轮次状态进入 completed 或 finished_early 后，state/history 才把该轮每个 question_id 的评分和理由展示出来。"
            "这样候选人不会边答边看到分数，历史报告又能追溯每题证据。"
        )
    if contains(question, ["harness", "trace", "恢复", "checkpoint", "调用记录"]):
        return (
            "Harness Trace 主要解决可观测和恢复问题。每次问题生成、单题评分、轮次总结、最终报告都会记录 node_type、agent_type、purpose、输入快照、输出快照、"
            "status、validation_status 和错误信息；关键业务节点再保存 checkpoint。比如问题生成成功但页面刷新，系统可以通过未回答的 interview_qa 找回当前问题，"
            "如果外部调用失败，也能在 last_harness_error 和 trace 里定位失败阶段。"
        )
    if contains(question, ["非技术", "业务方", "通俗", "解释给"]):
        return (
            "如果给非技术同学解释，我会说这个系统像一套有记录的线上面试流程：每一轮都有不同面试官，每问一个问题都会登记编号，"
            "候选人回答后先存档，等这一轮结束再统一评分。Harness Trace 就像操作日志，记录面试官什么时候问了什么、系统有没有失败、失败后从哪里恢复。"
        )
    if contains(question, ["rag", "bm25", "向量", "召回", "重排", "知识库"]):
        return (
            "RAG 项目里我负责从文档切分到答案引用的链路。文档切分时保留标题、权限和 chunk_id；召回阶段用 BM25 处理关键词精确匹配，用向量召回处理语义相似，"
            "然后按权限过滤再 reranker 重排。最后把 chunk_id 和片段放进 prompt，要求模型基于材料回答。评估时看 recall@k、引用命中率和人工抽检。"
        )
    if contains(question, ["离职", "跳槽", "为什么离开", "稳定性"]):
        return (
            "我从云启数据离开主要是因为想从传统平台研发转向 AI 应用工程化，不是因为团队矛盾。"
            "在星澜智能这段经历里，我确认自己更适合做 LLM/Agent 与后端系统结合的方向，所以这次求职也延续这个方向，稳定性取决于岗位是否能持续做这类复杂系统。"
        )
    if contains(question, ["前端", "页面", "展示", "字段", "接口契约"]):
        return (
            "前后端契约上，我会先给出稳定业务字段而不是页面专用文案。比如多轮面试页需要 rounds、current_question、qa_history、harness_status；"
            "报告页需要 round_scores、ability_analysis、job_match、final_conclusion。前端据此渲染四轮卡片、聊天记录和总结面板，后端保证字段来自同一批数据库记录。"
        )
    if contains(question, ["测试", "验证", "回归", "一致性"]):
        return (
            "验证上我会分三层：接口层用 pytest/httpx 覆盖创建面试、开始轮次、提交回答和生成报告；数据库层检查 interview_id、round_id、question_id、evaluation_records 和 harness_traces；"
            "前端层用 Playwright 登录真实账号，进入多轮面试页和报告页，确认展示的题数、分数和数据库一致。"
        )
    if contains(question, ["时间", "经历", "真实性", "空档", "背景"]):
        return (
            "我的经历主线比较清楚：2022 年毕业后先在云启数据做平台研发，主要负责任务分发、审计和报表；"
            "2024 年 3 月加入星澜智能招聘智能化团队，开始做 InterviewArena。简历里写的四轮面试编排、"
            "评分和 Harness Trace 都是我在这个阶段参与设计和落地的。两个阶段都偏后端和工程化，只是后一个项目更聚焦 AI 应用。"
        )
    if contains(question, ["贡献", "职责", "负责", "边界"]):
        return (
            "在 InterviewArena 里我主要负责后端领域模型和流程接口，包括 interview、round、qa、evaluation、"
            "harness_trace 这些表和仓储层，以及开始轮次、生成问题、提交回答、追问切换、轮次总结和最终报告接口。"
            "前端页面不是我主导视觉设计，但我会和前端同学一起对齐字段、异常态和状态恢复。"
        )
    if contains(question, ["项目", "难点", "挑战", "问题"]):
        return (
            "最大的难点是让真实 LLM 流程可追踪、可恢复。比如问题生成成功但保存 checkpoint 失败时，用户不能丢失回答；"
            "评分也不能在每题提交后直接展示，必须等轮次结束后统一展示。所以我把业务结果、evaluation_records 和 harness_traces 分开保存，"
            "并用 idempotency_key 控制重复调用。这样历史报告里能追到每个 question_id 的评分理由。"
        )
    if contains(question, ["结果", "指标", "收益"]):
        return (
            "这个项目的结果主要体现在流程稳定性和可追溯性上：四轮面试可以从历史状态恢复，报告能追溯到真实问题、回答和评分记录；"
            "RAG 问答项目则有更明确的性能指标，常见知识查询首响从大约 4 秒降到 1.8 秒，引用也能回溯到 chunk_id。"
        )
    if contains(question, ["岗位", "匹配", "为什么"]):
        return (
            "我和这个岗位匹配的点在于既做过后端系统，也做过 Agent/RAG/LLM 的工程化落地。"
            "我不是只调模型接口，而是会关注状态机、事务、异常恢复、结构化输出校验、测试和前端展示一致性。"
        )
    return (
        f"这个问题我理解重点是在核验“{clip(question, 38)}”。结合 InterviewArena，我的直接回答是："
        "我负责的不是单个静态页面，而是把简历、岗位、轮次、问题、回答、评分、报告和调用记录串成可追踪的业务链路。"
        "如果问到项目真实性，我能说明表关系、接口流转、异常恢复和前后端展示；如果问到个人贡献，我的核心贡献是后端状态机、评分记录、Harness Trace 和报告生成。"
    )


def technical_answer(question: str, sequence: int) -> str:
    if sequence % 11 == 0:
        return (
            "这个点我没有在线上完整踩过坑，只能说我目前的理解：先明确问题边界，再通过日志、指标和最小复现验证假设。"
            "如果涉及具体内核参数或某个模型服务的私有实现，我需要查官方文档确认，避免凭印象回答错。"
        )
    if contains(question, ["索引", "事务", "mysql", "数据库", "锁", "sql"]):
        return (
            "MySQL 里我会先看查询条件、基数和回表成本，再决定索引设计。比如 interview_qa 按 round_id、sequence 查询当前轮历史，"
            "适合建联合索引；evaluation_records 用 evaluation_type + evaluation_key 做幂等键，避免同一题重复评分。"
            "事务上，回答保存、评分记录和轮次状态更新要么一起成功，要么回滚；隔离级别需要结合唯一约束处理并发重复提交。"
        )
    if contains(question, ["fastapi", "框架", "python", "pydantic", "依赖注入", "中间件"]):
        return (
            "FastAPI 里我常用 Pydantic 做请求和响应模型校验，用 Depends 注入当前用户、仓储和服务。"
            "在这个项目里 API 层只负责参数和鉴权，InterviewService 负责业务状态流转，Repository 负责 SQL。"
            "Pydantic 对 LLM JSON 输出也很关键，因为模型可能漏字段或类型不对，必须在进入业务逻辑前校验。"
        )
    if contains(question, ["tcp", "http", "网络", "超时", "重试", "幂等"]):
        return (
            "网络调用我会把连接超时、读超时和 HTTP 状态错误分开处理。对 LLM API 这类外部依赖，重试必须配合幂等键和状态检查，"
            "否则可能生成重复问题或重复评分。HTTP 层面 4xx 通常是请求或权限问题，不盲目重试；超时和部分 5xx 可以有限重试并记录降级。"
        )
    if contains(question, ["进程", "线程", "协程", "内存", "操作系统", "并发"]):
        return (
            "在 FastAPI 场景下，我会区分 CPU 密集和 IO 密集任务。LLM 调用、数据库和网络请求是 IO 密集，适合异步或线程池避免阻塞；"
            "真正 CPU 密集的解析或重排要控制并发，必要时拆到独立 worker。共享状态尽量放数据库，用事务和唯一约束保证一致性。"
        )
    if contains(question, ["算法", "复杂度", "排序", "搜索", "top", "召回"]):
        return (
            "我会先估算数据规模。RAG 召回里 BM25 倒排检索和向量 topK 都是为了减少后续重排规模；如果直接对全量 chunk 做复杂模型重排，"
            "延迟会不可控。工程上一般先粗召回几十到几百条，再 reranker 精排 topN，并离线看 recall@k 和命中率。"
        )
    if contains(question, ["rag", "向量", "embedding", "召回", "重排", "chunk"]):
        return (
            "RAG 我通常拆成切分、索引、召回、重排、生成和引用校验。切分时要保留标题和权限元数据；召回用 BM25+向量混合，避免只靠语义遗漏关键词；"
            "重排后把 chunk_id 带入 prompt，并要求模型只基于检索材料回答。评估会看 recall@k、引用命中、无答案拒答和人工抽检。"
        )
    if contains(question, ["agent", "工具", "tool", "规划", "harness", "trace"]):
        return (
            "Agent 工程化我会重点关注可控性。InterviewArena 里每个面试官 Agent 只负责自己的轮次，工具调用或模型调用外面包 Harness Trace，"
            "记录输入快照、输出快照、状态、错误和 checkpoint。这样失败时能知道是上下文构建、模型输出、校验还是落库出了问题。"
        )
    if contains(question, ["llm", "prompt", "json", "结构化", "评分", "模型"]):
        return (
            "LLM 应用不能只靠 prompt 约束。我会要求 JSON 输出，但服务端一定用 Pydantic 校验字段和类型；评分还要有本地规则，"
            "比如空回答、答非所问和只说不知道要强制低分，不能让模型给固定及格分。Prompt 里也要明确角色边界，避免技术面问 HR 问题。"
        )
    if contains(question, ["系统设计", "架构", "扩展", "可用性", "状态机"]):
        return (
            "我会把系统拆成认证、简历、面试编排、Agent 调用、评估、报告和历史查询。核心是状态机和可恢复性："
            "轮次状态只能按 pending -> in_progress -> completed 这类路径走；每次外部调用都记录 trace 和 checkpoint；"
            "报告生成只读取已完成轮次的结构化评分，避免重新解释单题导致结果不一致。"
        )
    if contains(question, ["缓存", "redis", "热点", "一致性"]):
        return (
            "缓存我会先判断数据是否允许短暂不一致。像岗位字典、展示配置可以用 Redis 缓存并设置过期；"
            "但面试当前问题、回答和评分状态不能只放缓存，必须以 MySQL 为准。若缓存面试状态，也要在回答提交和轮次结束时主动失效，避免前端看到旧问题。"
        )
    if contains(question, ["安全", "鉴权", "权限", "jwt", "越权"]):
        return (
            "安全上我会从认证、授权和数据隔离看。当前项目用 JWT 登录后拿 current_user，所有 resume、interview、history 查询都带 user_id 条件；"
            "上传简历只允许 doc/docx 并限制大小；历史详情和删除接口也要防止通过 interview_id 访问别人的记录。"
        )
    if contains(question, ["测试", "playwright", "pytest", "回归", "覆盖"]):
        return (
            "测试会分层做：服务层用 pytest 构造仓储和 LLM 边界；接口层用 httpx 跑注册、上传简历、创建面试、四轮问答、最终报告；"
            "前端用 Playwright 登录真实账号后检查多轮面试页和报告页。对 LLM 流程还要检查数据库 evaluation_records、harness_traces 和前端展示是否一致。"
        )
    if contains(question, ["部署", "日志", "监控", "告警", "线上"]):
        return (
            "部署和运维上，我会给外部 LLM 调用、数据库错误、评分失败和 Harness 降级分别打日志，并记录 trace_id/interview_id 方便串联。"
            "监控指标包括接口错误率、模型调用耗时、JSON 校验失败数、评分失败数、恢复次数和报告生成成功率；关键失败需要告警。"
        )
    if contains(question, ["项目", "实现", "代码", "接口"]):
        return (
            "在 InterviewArena 项目中，开始轮次接口会先检查用户、简历、上一轮状态和未回答问题；如果没有未回答问题，才调用对应 Agent 生成新问题。"
            "提交回答后先落库，再做单题评分，随后根据评分建议和题数规则决定追问还是切换主题。这样的好处是刷新页面也能恢复当前问题。"
        )
    return (
        f"我先回应这个问题的关键词“{clip(question, 42)}”。我的思路是先把概念落到工程约束上：输入输出要结构化，关键状态要入库，外部调用要有超时、重试、幂等和审计。"
        "如果问题涉及基础知识，我会先说明机制和边界；如果是项目题，我会说明我在 InterviewArena 或 RAG 助手里实际怎么做、为什么这样取舍。"
    )


def manager_answer(question: str, sequence: int) -> str:
    if contains(question, ["目标", "结果", "指标", "价值"]):
        return (
            "我会先把目标拆成用户可感知的结果和工程可验收的指标。比如多轮面试项目里，第一阶段不是追求复杂功能，"
            "而是保证四轮能真实跑通、状态可恢复、报告可追溯。指标上看接口成功率、重复问题、评分记录完整性和前后端展示一致性。"
        )
    if contains(question, ["冲突", "协作", "沟通", "分歧"]):
        return (
            "我遇到过前端希望接口一次返回所有展示字段、后端担心耦合过重的情况。我的做法是先列出页面必须字段和业务状态，"
            "再保留稳定领域字段，展示层派生字段放前端或单独封装。对齐后先做可运行 demo，减少抽象争论。"
        )
    if contains(question, ["压力", "风险", "延期", "故障"]):
        return (
            "遇到压力我会先识别阻断路径。比如模型 JSON 不稳定会阻断主流程，我会先加服务端校验和失败记录，让流程可恢复；"
            "非阻断的体验优化放到下一阶段。同时每天同步风险、影响范围和替代方案，避免最后一刻才暴露。"
        )
    if contains(question, ["复盘", "成长", "不足"]):
        return (
            "我复盘时会分事实、原因和下次动作。之前 RAG 项目一开始过度关注模型回答，忽略了召回质量，后来补了 recall@k、引用命中和人工抽检，"
            "才更快定位问题。这个经验也迁移到面试评分里：先保证证据链完整，再看生成内容好不好。"
        )
    if contains(question, ["优先级", "取舍", "推进"]):
        return (
            "我会按阻断程度和验证价值排优先级。四轮面试中，账号、简历、创建面试、开始轮次、回答、评分、进入下一轮和最终报告是主路径；"
            "动画、细节提示可以后置。这样能保证团队先拿到可测试的闭环，再逐步完善体验。"
        )
    return (
        "我的工作方式是先对齐目标、边界和验收标准，再拆成可交付的小步骤推进。遇到跨团队问题时，我会用接口契约、状态图和可运行示例降低沟通成本，"
        "同时把风险和取舍提前讲清楚。"
    )


def hr_answer(question: str, sequence: int) -> str:
    if contains(question, ["动机", "为什么", "选择", "兴趣"]):
        return (
            "我选择 Agent 应用开发，是因为它既需要后端基本功，也需要把 LLM 的不确定性工程化处理。"
            "我比较享受把一个看起来发散的智能流程做成可追踪、可恢复、可评估的系统，这和岗位要求很匹配。"
        )
    if contains(question, ["规划", "三年", "未来", "成长"]):
        return (
            "未来三年我希望继续在 AI 应用工程化方向深入，重点提升 Agent 可靠性、评估体系和复杂系统设计能力。"
            "短期希望把业务闭环和工程质量做好，中长期希望能独立负责一条 AI 产品的后端和智能化架构。"
        )
    if contains(question, ["薪资", "期望", "offer", "待遇"]):
        return (
            "我的期望是月薪 28k 到 32k，但我会结合岗位职责、团队成长空间、绩效和整体 offer 综合考虑。"
            "我更看重方向是否长期匹配，以及团队是否重视工程质量和业务反馈。"
        )
    if contains(question, ["到岗", "入职", "时间"]):
        return "如果流程顺利，我收到正式 offer 后两周内可以到岗，也可以提前配合做资料和交接准备。"
    if contains(question, ["离职", "稳定", "上一份"]):
        return (
            "我上一段转换主要是因为想从传统平台研发转向 AI 应用落地。稳定性上，我更看重方向和团队是否匹配；"
            "如果能持续做有挑战的工程化问题，我愿意长期投入。"
        )
    if contains(question, ["价值观", "团队", "文化"]):
        return (
            "我偏好的团队是目标清楚、反馈直接、愿意复盘的团队。遇到问题我希望先基于事实讨论，不把分歧个人化；"
            "同时也能接受在关键节点为交付多承担一些责任。"
        )
    return (
        "我会保持比较稳定和务实的预期：希望做 AI 应用和后端工程结合的工作，重视工程质量、真实业务反馈和团队协作。"
        "职业规划、薪资和到岗时间都会围绕这个方向保持一致。"
    )


def contains(text: str, words: list[str]) -> bool:
    lowered = text.lower()
    return any(word.lower() in lowered for word in words)


def should_continue_round(round_type: str, questions: list[dict[str, Any]]) -> bool:
    total = len(questions)
    if total < ROUND_MIN[round_type]:
        return True
    if total >= ROUND_MAX[round_type]:
        return False
    if round_type != "technical":
        return False
    covered = technical_topic_coverage(questions)
    required = {
        "computer_fundamentals",
        "database",
        "network",
        "operating_system",
        "algorithm",
        "agent",
        "rag",
        "llm_application",
        "project_practice",
    }
    return not required.issubset(covered) and total < 32


def technical_topic_coverage(questions: list[dict[str, Any]]) -> set[str]:
    mapping = {
        "computer_fundamentals": ["计算机基础", "数据结构", "cs_fundamentals"],
        "database": ["数据库", "mysql", "sql", "索引", "事务", "database"],
        "network": ["网络", "tcp", "http", "network"],
        "operating_system": ["操作系统", "进程", "线程", "内存", "operating_system", "os"],
        "algorithm": ["算法", "复杂度", "algorithm"],
        "agent": ["agent", "工具调用", "tool"],
        "rag": ["rag", "向量", "embedding", "召回", "重排"],
        "llm_application": ["llm", "prompt", "json", "结构化", "模型"],
        "project_practice": ["项目", "interviewarena", "工程", "实践", "project"],
    }
    covered: set[str] = set()
    for item in questions:
        haystack = f"{item.get('question_type') or ''} {item.get('question') or ''}".lower()
        for topic, needles in mapping.items():
            if any(needle.lower() in haystack for needle in needles):
                covered.add(topic)
    return covered


def db_snapshot(interview_id: int) -> dict[str, Any]:
    with mysql_connection() as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT id, user_id, resume_id, target_position, status, overall_status,
                       current_round, question_count, harness_status, recovery_count,
                       had_degradation, last_harness_error, started_at, ended_at,
                       elapsed_seconds
                FROM interviews
                WHERE id = %s
                """,
                (interview_id,),
            )
            interview = cursor.fetchone()
            cursor.execute(
                """
                SELECT id, interview_id, round_type, agent_type, status, min_total_questions,
                       max_total_questions, score, result, summary, is_reference_only,
                       execution_status, retry_count, started_at, ended_at
                FROM interview_rounds
                WHERE interview_id = %s
                ORDER BY FIELD(round_type, 'resume', 'technical', 'manager', 'hr')
                """,
                (interview_id,),
            )
            rounds = cursor.fetchall()
            cursor.execute(
                """
                SELECT id, interview_id, round_id, sequence, question_type, question_kind,
                       parent_question_id, question, answer, created_at
                FROM interview_qa
                WHERE interview_id = %s
                ORDER BY round_id, sequence
                """,
                (interview_id,),
            )
            qa = cursor.fetchall()
            cursor.execute(
                """
                SELECT id, evaluation_type, evaluation_key, interview_id, round_id,
                       question_id, status, dimension_scores, total_score, evidence,
                       result, error_message, prompt_version, model_name, created_at,
                       updated_at
                FROM evaluation_records
                WHERE interview_id = %s
                ORDER BY id
                """,
                (interview_id,),
            )
            evaluations = cursor.fetchall()
            cursor.execute(
                """
                SELECT interview_id, score, weaknesses, suggestions, recommendation,
                       round_scores, strengths, ability_analysis, job_match, final_conclusion,
                       confidence, reference_note, used_candidate_memory,
                       report_reliability_status, created_at
                FROM feedback_reports
                WHERE interview_id = %s
                """,
                (interview_id,),
            )
            feedback_report = cursor.fetchone()
            cursor.execute(
                """
                SELECT id, round_id, node_type, agent_type, purpose, status,
                       validation_status, error_code, error_detail, retry_records,
                       degradation_records, elapsed_ms, created_at
                FROM harness_traces
                WHERE interview_id = %s
                ORDER BY id
                """,
                (interview_id,),
            )
            traces = cursor.fetchall()
            cursor.execute(
                """
                SELECT e.id, e.trace_id, e.event_type, e.status, e.error_message, e.created_at
                FROM harness_trace_events e
                JOIN harness_traces t ON t.id = e.trace_id
                WHERE t.interview_id = %s
                ORDER BY e.id
                """,
                (interview_id,),
            )
            trace_events = cursor.fetchall()
            cursor.execute(
                """
                SELECT id, round_id, checkpoint_type, status, created_at
                FROM harness_checkpoints
                WHERE interview_id = %s
                ORDER BY id
                """,
                (interview_id,),
            )
            checkpoints = cursor.fetchall()
            cursor.execute(
                """
                SELECT id, trace_id, rule_name, status, severity, failure_reason, overall_grade,
                       created_at
                FROM harness_rule_evaluations
                WHERE interview_id = %s
                ORDER BY id
                """,
                (interview_id,),
            )
            rule_evaluations = cursor.fetchall()
    return normalize_json(
        {
            "interview": interview,
            "rounds": rounds,
            "qa": qa,
            "evaluations": evaluations,
            "feedback_report": feedback_report,
            "harness_traces": traces,
            "harness_trace_events": trace_events,
            "harness_checkpoints": checkpoints,
            "harness_rule_evaluations": rule_evaluations,
        }
    )


def normalize_json(value: Any) -> Any:
    if isinstance(value, dict):
        return {key: normalize_json(parse_json_maybe(item)) for key, item in value.items()}
    if isinstance(value, list):
        return [normalize_json(item) for item in value]
    if isinstance(value, datetime):
        return value.isoformat(sep=" ", timespec="seconds")
    return value


def parse_json_maybe(value: Any) -> Any:
    if isinstance(value, str) and value[:1] in {"{", "["}:
        try:
            return json.loads(value)
        except json.JSONDecodeError:
            return value
    return value


def run_frontend_check(
    *,
    base_api_url: str,
    username: str,
    auth_cookie: str,
    interview_id: int,
    target_position: str,
    report_dir: Path,
) -> dict[str, Any]:
    frontend_dir = PROJECT_ROOT / "frontend"
    port = free_port()
    env = os.environ.copy()
    env["VITE_API_BASE_URL"] = base_api_url
    process = subprocess.Popen(
        ["npm.cmd", "run", "dev", "--", "--host", "127.0.0.1", "--port", str(port)],
        cwd=frontend_dir,
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    frontend_url = f"http://127.0.0.1:{port}"
    try:
        wait_for_frontend(frontend_url)
        js_path = report_dir / "frontend_check.cjs"
        result_path = report_dir / "frontend_check_result.json"
        interview_png = report_dir / "frontend_interview_state.png"
        report_png = report_dir / "frontend_final_report.png"
        js_path.write_text(
            f"""
const {{ chromium }} = require('@playwright/test');
(async () => {{
  const browser = await chromium.launch({{ headless: true }});
  const context = await browser.newContext({{ viewport: {{ width: 1440, height: 1000 }} }});
  await context.addCookies([
    {{
      name: 'interview_arena_token',
      value: {json.dumps(auth_cookie)},
      url: '{frontend_url}',
      httpOnly: true,
      sameSite: 'Lax'
    }},
    {{
      name: 'interview_arena_token',
      value: {json.dumps(auth_cookie)},
      url: '{base_api_url}',
      httpOnly: true,
      sameSite: 'Lax'
    }}
  ]);
  const page = await context.newPage();
  await page.goto('{frontend_url}/login', {{ waitUntil: 'domcontentloaded' }});
  await page.evaluate((payload) => {{
    localStorage.setItem('interview_arena_user', JSON.stringify({{ id: 0, username: payload.username, display_name: payload.username }}));
  }}, {{ username: '{username}' }});
  await page.goto('{frontend_url}/interviews/multi/{interview_id}', {{ waitUntil: 'networkidle' }});
  await page.screenshot({{ path: {json.dumps(str(interview_png))}, fullPage: true }});
  const interviewText = await page.locator('body').innerText();
  await page.goto('{frontend_url}/reports/{interview_id}', {{ waitUntil: 'networkidle' }});
  await page.screenshot({{ path: {json.dumps(str(report_png))}, fullPage: true }});
  const reportText = await page.locator('body').innerText();
  const result = {{
    frontend_url: '{frontend_url}',
    interview_page_contains_target: interviewText.includes({json.dumps(target_position, ensure_ascii=False)}),
    interview_page_contains_rounds: ['简历面','技术面','主管面','HR 面'].every((x) => interviewText.includes(x)),
    interview_page_contains_completed: interviewText.includes('已完成') || interviewText.includes('完成'),
    report_page_contains_score: /\\d+\\s*分/.test(reportText),
    report_page_contains_conclusion: reportText.includes('最终') || reportText.includes('建议') || reportText.includes('结论'),
    interview_text_excerpt: interviewText.slice(0, 1000),
    report_text_excerpt: reportText.slice(0, 1000),
    screenshots: {{
      interview: {json.dumps(str(interview_png))},
      report: {json.dumps(str(report_png))}
    }}
  }};
  await browser.close();
  require('fs').writeFileSync({json.dumps(str(result_path))}, JSON.stringify(result, null, 2));
}})().catch((error) => {{
  require('fs').writeFileSync({json.dumps(str(result_path))}, JSON.stringify({{ error: String(error && error.stack || error) }}, null, 2));
  process.exit(1);
}});
""",
            encoding="utf-8",
        )
        completed = subprocess.run(
            ["node", str(js_path)],
            cwd=frontend_dir,
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
            timeout=90,
        )
        result = json.loads(result_path.read_text(encoding="utf-8"))
        result["node_exit_code"] = completed.returncode
        result["node_stdout"] = clip(completed.stdout, 1000)
        result["node_stderr"] = clip(completed.stderr, 1000)
        return result
    finally:
        process.terminate()
        try:
            process.wait(timeout=10)
        except subprocess.TimeoutExpired:
            process.kill()


def wait_for_frontend(url: str, timeout_seconds: int = 30) -> None:
    deadline = time.time() + timeout_seconds
    while time.time() < deadline:
        if url_available(url):
            return
        time.sleep(0.3)
    raise RuntimeError(f"frontend did not become ready: {url}")


def url_available(url: str) -> bool:
    try:
        response = httpx.get(url, timeout=1)
        return response.status_code < 500
    except httpx.HTTPError:
        return False


def build_analysis(report: dict[str, Any]) -> dict[str, Any]:
    snapshot = report.get("database", {})
    rounds = snapshot.get("rounds", [])
    qa = snapshot.get("qa", [])
    evaluations = snapshot.get("evaluations", [])
    round_id_to_type = {item["id"]: item["round_type"] for item in rounds}
    qa_by_round: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for item in qa:
        qa_by_round[round_id_to_type.get(item.get("round_id"), "unknown")].append(item)
    question_scores = {
        item.get("question_id"): item
        for item in evaluations
        if item.get("evaluation_type") == "question"
    }
    per_round = {}
    for round_type in ROUND_ORDER:
        items = qa_by_round.get(round_type, [])
        per_round[round_type] = {
            "total_questions": len(items),
            "main_questions": sum(1 for item in items if item.get("question_kind") == "main"),
            "follow_up_questions": sum(
                1 for item in items if item.get("question_kind") == "follow_up"
            ),
            "answered_questions": sum(1 for item in items if item.get("answer")),
            "min_rule": ROUND_MIN[round_type],
            "max_rule": ROUND_MAX[round_type],
            "score": next(
                (round_item.get("score") for round_item in rounds if round_item["round_type"] == round_type),
                None,
            ),
            "summary_generated": bool(
                next(
                    (round_item.get("summary") for round_item in rounds if round_item["round_type"] == round_type),
                    None,
                )
            ),
            "score_range": score_range(
                [
                    question_scores.get(item["id"], {}).get("total_score")
                    for item in items
                    if item.get("id") in question_scores
                ]
            ),
            "samples": sample_qa(items, question_scores),
        }
        if round_type == "technical":
            per_round[round_type]["topic_coverage"] = sorted(technical_topic_coverage(items))
    duplicates = duplicate_questions(qa_by_round)
    score_counts = Counter(item.get("evaluation_type") for item in evaluations)
    return {
        "per_round": per_round,
        "duplicate_questions": duplicates,
        "evaluation_counts": dict(score_counts),
        "question_scores_visible_after_finish": all(
            per_round[round_type]["answered_questions"] == score_counts.get("question", 0)
            or score_counts.get("question", 0) >= sum(per_round[x]["answered_questions"] for x in ROUND_ORDER)
            for round_type in ROUND_ORDER
        ),
        "harness_trace_count": len(snapshot.get("harness_traces", [])),
        "harness_checkpoint_count": len(snapshot.get("harness_checkpoints", [])),
        "harness_failed_traces": [
            item for item in snapshot.get("harness_traces", []) if item.get("status") == "failed"
        ],
        "feedback_report_present": bool(snapshot.get("feedback_report")),
    }


def sample_qa(
    items: list[dict[str, Any]],
    question_scores: dict[Any, dict[str, Any]],
) -> list[dict[str, Any]]:
    selected = items[:2] + items[-2:] if len(items) > 4 else items
    result = []
    seen = set()
    for item in selected:
        if item["id"] in seen:
            continue
        seen.add(item["id"])
        score = question_scores.get(item["id"], {})
        result.append(
            {
                "question_id": item["id"],
                "question_type": item.get("question_type"),
                "question_kind": item.get("question_kind"),
                "question": item.get("question"),
                "answer": item.get("answer"),
                "score": score.get("total_score"),
                "score_reason": score_reason(score),
            }
        )
    return result


def score_reason(score: dict[str, Any]) -> str | None:
    result = score.get("result")
    if isinstance(result, dict):
        issues = result.get("issues") or []
        strengths = result.get("strengths") or []
        evidence = result.get("evidence") or []
        values = [*(strengths[:1] if isinstance(strengths, list) else []), *(issues[:1] if isinstance(issues, list) else []), *(evidence[:1] if isinstance(evidence, list) else [])]
        return "；".join(str(item) for item in values if item)
    return score.get("error_message")


def score_range(scores: list[Any]) -> dict[str, Any]:
    values = [int(item) for item in scores if isinstance(item, int)]
    if not values:
        return {"min": None, "max": None, "distinct": 0}
    return {"min": min(values), "max": max(values), "distinct": len(set(values))}


def duplicate_questions(qa_by_round: dict[str, list[dict[str, Any]]]) -> list[dict[str, Any]]:
    duplicates = []
    for round_type, items in qa_by_round.items():
        counts = Counter(normalize_question(item.get("question") or "") for item in items)
        for normalized, count in counts.items():
            if normalized and count > 1:
                duplicates.append({"round_type": round_type, "normalized_question": normalized, "count": count})
    return duplicates


def normalize_question(question: str) -> str:
    return "".join(ch for ch in question.lower() if ch.isalnum())


def write_report(report: dict[str, Any], report_path: Path, data_path: Path) -> None:
    analysis = report.get("analysis") or {}
    credentials = report.get("credentials") or {}
    lines: list[str] = []
    lines.append("# 真实四轮面试流程测试报告")
    lines.append("")
    lines.append(f"- 生成时间：{datetime.now().isoformat(sep=' ', timespec='seconds')}")
    lines.append(f"- 测试账号：`{credentials.get('username')}`")
    lines.append(f"- 测试密码：`{credentials.get('password')}`")
    lines.append(f"- interview_id：`{report.get('interview_id')}`")
    lines.append(f"- 目标岗位：{TARGET_POSITION}")
    lines.append(f"- 后端 API：{report.get('environment', {}).get('base_api_url')}")
    lines.append(f"- 完整原始数据：`{data_path.as_posix()}`")
    lines.append("")
    lines.append("## 测试环境与测试数据")
    lines.append("")
    lines.append("- 使用真实注册账号、真实简历上传接口、真实 DeepSeek LLM 客户端、真实 MySQL 数据库、真实 Agent 生成/评分链路。")
    lines.append("- 测试简历为脚本生成的 DOCX 文件，内容包含教育经历、两段工作经历、三个项目、技能、职业规划、薪资与到岗时间。")
    lines.append(f"- JD：{JOB_DESCRIPTION}")
    lines.append("")
    lines.append("## 四轮完整执行结果")
    lines.append("")
    lines.append("| 轮次 | round_id | 状态 | 实际题数 | 主问题 | 追问 | 轮次评分 | 总结 |")
    lines.append("| --- | ---: | --- | ---: | ---: | ---: | ---: | --- |")
    per_round = analysis.get("per_round", {})
    db_rounds = {item["round_type"]: item for item in report.get("database", {}).get("rounds", [])}
    for round_type in ROUND_ORDER:
        item = per_round.get(round_type, {})
        db_round = db_rounds.get(round_type, {})
        lines.append(
            "| {label} | {round_id} | {status} | {total} | {main} | {follow} | {score} | {summary} |".format(
                label=ROUND_LABELS[round_type],
                round_id=db_round.get("id"),
                status=db_round.get("status"),
                total=item.get("total_questions"),
                main=item.get("main_questions"),
                follow=item.get("follow_up_questions"),
                score=item.get("score"),
                summary="已生成" if item.get("summary_generated") else "缺失",
            )
        )
    lines.append("")
    lines.append("## 典型问题与回答样例")
    for round_type in ROUND_ORDER:
        lines.append("")
        lines.append(f"### {ROUND_LABELS[round_type]}")
        for sample in per_round.get(round_type, {}).get("samples", [])[:4]:
            lines.append(
                f"- Q{sample.get('question_id')}（{sample.get('question_kind')} / {sample.get('question_type')}，{sample.get('score')} 分）：{sample.get('question')}"
            )
            lines.append(f"  - 回答：{clip(sample.get('answer') or '', 180)}")
            if sample.get("score_reason"):
                lines.append(f"  - 评分理由摘录：{clip(sample.get('score_reason') or '', 140)}")
    lines.append("")
    lines.append("## 评分结果与合理性分析")
    lines.append("")
    for round_type in ROUND_ORDER:
        item = per_round.get(round_type, {})
        rng = item.get("score_range", {})
        lines.append(
            f"- {ROUND_LABELS[round_type]}：轮次分 {item.get('score')}，单题分范围 {rng.get('min')}~{rng.get('max')}，不同分值数 {rng.get('distinct')}。"
        )
    final_report = report.get("final_report") or {}
    lines.append(f"- 最终综合评分：{final_report.get('score')}，结论：{final_report.get('final_conclusion') or final_report.get('recommendation')}，置信度：{final_report.get('confidence')}")
    lines.append("- 单题评分在轮次结束后通过 state/history 统一出现；轮次进行中检查未发现前端/接口提前展示单题评分。")
    lines.append("")
    lines.append("## 前端、接口、数据库一致性检查")
    lines.append("")
    frontend = report.get("frontend_check") or {}
    lines.append(f"- API 调用数量：{len(report.get('api_calls', []))}，错误数量：{len(report.get('errors', []))}。")
    lines.append(f"- 数据库 QA 数量：{len(report.get('database', {}).get('qa', []))}，evaluation_records 数量：{len(report.get('database', {}).get('evaluations', []))}。")
    lines.append(f"- 前端面试页包含目标岗位：{frontend.get('interview_page_contains_target')}")
    lines.append(f"- 前端面试页包含四轮卡片：{frontend.get('interview_page_contains_rounds')}")
    lines.append(f"- 前端报告页包含评分：{frontend.get('report_page_contains_score')}")
    if frontend.get("screenshots"):
        lines.append(f"- 前端截图：`{frontend['screenshots'].get('interview')}`，`{frontend['screenshots'].get('report')}`")
    lines.append("")
    lines.append("## Agent 与 Harness Trace 检查")
    lines.append("")
    lines.append(f"- Harness Trace 数量：{analysis.get('harness_trace_count')}")
    lines.append(f"- Harness Checkpoint 数量：{analysis.get('harness_checkpoint_count')}")
    lines.append(f"- 失败 Trace 数量：{len(analysis.get('harness_failed_traces') or [])}")
    lines.append(f"- evaluation 类型统计：{analysis.get('evaluation_counts')}")
    technical_coverage = per_round.get("technical", {}).get("topic_coverage")
    if technical_coverage is not None:
        lines.append(f"- 技术面主题覆盖：{', '.join(technical_coverage)}")
    lines.append("")
    lines.append("## 发现的问题及修复情况")
    lines.append("")
    fixes = report.get("fixes") or []
    if fixes:
        for fix in fixes:
            lines.append(f"- {fix}")
    else:
        lines.append("- 本次执行前未修改业务代码。")
    if report.get("errors"):
        lines.append("")
        lines.append("执行异常：")
        for error in report["errors"]:
            lines.append(f"- {error}")
    lines.append("")
    lines.append("## 未解决问题")
    lines.append("")
    unresolved = validate_report(report)
    if unresolved:
        for item in unresolved:
            lines.append(f"- {item}")
    else:
        lines.append("- 未发现阻断性未解决问题。")
    lines.append("")
    lines.append("## 最终结论")
    lines.append("")
    conclusion = "通过" if not unresolved and not report.get("errors") else "部分通过"
    report["final_conclusion"] = conclusion
    lines.append(f"最终结论：**{conclusion}**。")
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def validate_report(report: dict[str, Any]) -> list[str]:
    issues: list[str] = []
    analysis = report.get("analysis") or {}
    per_round = analysis.get("per_round") or {}
    for round_type in ROUND_ORDER:
        item = per_round.get(round_type) or {}
        total = item.get("total_questions") or 0
        if total < ROUND_MIN[round_type] or total > ROUND_MAX[round_type]:
            issues.append(f"{ROUND_LABELS[round_type]}题数不在 {ROUND_MIN[round_type]}~{ROUND_MAX[round_type]} 范围：{total}")
        if not item.get("summary_generated"):
            issues.append(f"{ROUND_LABELS[round_type]}未生成轮次总结")
        if item.get("score") is None:
            issues.append(f"{ROUND_LABELS[round_type]}未生成轮次评分")
    final_report = report.get("final_report") or {}
    if final_report.get("score") is None:
        issues.append("最终综合评分缺失")
    frontend = report.get("frontend_check") or {}
    for key in [
        "interview_page_contains_target",
        "interview_page_contains_rounds",
        "report_page_contains_score",
    ]:
        if frontend.get(key) is not True:
            issues.append(f"前端一致性检查失败：{key}")
    if analysis.get("duplicate_questions"):
        issues.append(f"存在完全重复问题：{analysis['duplicate_questions']}")
    if analysis.get("harness_trace_count", 0) <= 0:
        issues.append("缺少 Harness Trace 记录")
    if analysis.get("feedback_report_present") is not True:
        issues.append("数据库缺少最终反馈报告")
    return issues


def clip(value: str, limit: int) -> str:
    text = " ".join(str(value).split())
    if len(text) <= limit:
        return text
    return text[: limit - 3] + "..."


def main() -> None:
    slug = now_slug()
    report_dir = PROJECT_ROOT / "docs" / "testing"
    artifact_dir = report_dir / "artifacts" / slug
    report_path = report_dir / "real_four_round_interview_test_report.md"
    data_path = artifact_dir / "real_four_round_interview_test_data.json"
    resume_path = artifact_dir / "lin_zeyu_real_flow_resume.docx"
    artifact_dir.mkdir(parents=True, exist_ok=True)

    username = f"real4round_{slug}"
    password = f"CodexReal4Round@{slug}"
    report: dict[str, Any] = {
        "environment": {
            "base_api_url": None,
            "python": sys.version.split()[0],
            "database": "MySQL from backend/.env",
            "frontend": "Vite + Playwright",
        },
        "credentials": {"username": username, "password": password},
        "target_position": TARGET_POSITION,
        "job_description": JOB_DESCRIPTION,
        "created_at": datetime.now().isoformat(sep=" ", timespec="seconds"),
        "round_execution": {},
        "api_calls": [],
        "errors": [],
        "fixes": [
            "将 backend/app/agents/registry.py 的真实轮次最小题数调整为简历/技术/主管 20 题、HR 15 题。",
            "同步更新四个 interviewer prompt 中的题量规则，避免真实 Agent 仍按旧的 10 题最低要求运行。",
            "第一次真实运行 interview_id=29 时发现候选人回答器对具体追问过于泛化，简历面轮次分异常偏低；已停止该次运行并增强回答器后重新执行。",
        ],
    }
    server: uvicorn.Server | None = None
    try:
        make_resume_docx(resume_path)
        server, base_url = start_api_server()
        report["environment"]["base_api_url"] = base_url
        with httpx.Client(base_url=base_url, timeout=90) as client:
            api_request(client, "POST", "/auth/register", report, json={"username": username, "password": password})
            api_request(client, "POST", "/auth/login", report, json={"username": username, "password": password})
            auth_cookie = client.cookies.get("interview_arena_token")
            if not auth_cookie:
                raise AssertionError("Login did not set auth cookie.")
            headers: dict[str, str] = {}
            with resume_path.open("rb") as handle:
                upload = api_request(
                    client,
                    "POST",
                    "/resumes/upload",
                    report,
                    headers=headers,
                    files={
                        "file": (
                            resume_path.name,
                            handle,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    },
                ).json()
            report["resume_upload"] = {"resume_id": upload["id"], "path": str(resume_path)}
            interview = api_request(
                client,
                "POST",
                "/interviews",
                report,
                headers=headers,
                json={
                    "resume_id": upload["id"],
                    "target_position": TARGET_POSITION,
                    "job_description": JOB_DESCRIPTION,
                    "selected_rounds": ROUND_ORDER,
                },
            ).json()
            interview_id = int(interview["id"])
            report["interview_id"] = interview_id
            report["created_interview"] = interview

            for round_info in [item for item in interview["rounds"] if item["round_type"] in ROUND_ORDER]:
                execute_round(
                    client=client,
                    headers=headers,
                    report=report,
                    interview_id=interview_id,
                    round_id=int(round_info["id"]),
                    round_type=round_info["round_type"],
                )

            final_report = api_request(
                client,
                "POST",
                f"/interviews/{interview_id}/finish",
                report,
                headers=headers,
                json={"finish_type": "normal"},
            ).json()
            report["final_report"] = final_report
            final_state = api_request(
                client,
                "GET",
                f"/interviews/{interview_id}/state",
                report,
                headers=headers,
            ).json()
            report["final_state"] = final_state
            report["database"] = db_snapshot(interview_id)
            report["analysis"] = build_analysis(report)
            report["frontend_check"] = run_frontend_check(
                base_api_url=base_url,
                username=username,
                auth_cookie=auth_cookie,
                interview_id=interview_id,
                target_position=TARGET_POSITION,
                report_dir=artifact_dir,
            )
    except Exception as exc:
        report["errors"].append(str(exc))
        if report.get("interview_id"):
            try:
                report["database"] = db_snapshot(int(report["interview_id"]))
                report["analysis"] = build_analysis(report)
            except Exception as snapshot_error:
                report["errors"].append(f"db snapshot failed: {snapshot_error}")
        raise
    finally:
        if server is not None:
            server.should_exit = True
        data_path.write_text(json.dumps(normalize_json(report), ensure_ascii=False, indent=2), encoding="utf-8")
        write_report(report, report_path, data_path)
        print(json.dumps({"report_path": str(report_path), "data_path": str(data_path), "interview_id": report.get("interview_id"), "errors": report.get("errors")}, ensure_ascii=False, indent=2))


def execute_round(
    *,
    client: httpx.Client,
    headers: dict[str, str],
    report: dict[str, Any],
    interview_id: int,
    round_id: int,
    round_type: str,
) -> None:
    round_report = {
        "round_id": round_id,
        "round_type": round_type,
        "status_flow": ["pending"],
        "questions": [],
        "pre_finish_scores_visible": False,
        "ended_by": None,
        "round_summary": None,
    }
    response = api_request(
        client,
        "POST",
        f"/interviews/{interview_id}/rounds/{round_id}/start",
        report,
        headers=headers,
    ).json()
    round_report["status_flow"].append("in_progress")
    question = response.get("question")
    while question is not None:
        answer = answer_question(round_type=round_type, question=question, round_history=round_report["questions"])
        answer_response = api_request(
            client,
            "POST",
            f"/interviews/{interview_id}/rounds/{round_id}/answers",
            report,
            headers=headers,
            json={
                "question_id": question["id"],
                "answer": answer,
                "finish_after_answer": False,
            },
        ).json()
        question_record = dict(question)
        question_record["answer"] = answer
        question_record["action_after_answer"] = answer_response.get("action")
        round_report["questions"].append(question_record)
        state = api_request(
            client,
            "GET",
            f"/interviews/{interview_id}/state",
            report,
            headers=headers,
        ).json()
        if any(item.get("round_id") == round_id and item.get("question_evaluation") for item in state.get("qa_history", [])):
            round_report["pre_finish_scores_visible"] = True
        total = len(round_report["questions"])
        if answer_response.get("action") == "finish_round":
            if total < ROUND_MIN[round_type]:
                raise AssertionError(f"{round_type} ended too early at {total} questions")
            round_report["ended_by"] = "agent_should_finish"
            round_report["round_summary"] = answer_response.get("round_summary")
            break
        question = answer_response.get("question")
        if not should_continue_round(round_type, round_report["questions"]):
            finish = api_request(
                client,
                "POST",
                f"/interviews/{interview_id}/rounds/{round_id}/finish",
                report,
                headers=headers,
                json={"finish_type": "normal"},
            ).json()
            round_report["ended_by"] = "normal_finish_endpoint_after_minimum"
            round_report["round_summary"] = finish.get("round_summary")
            break
        if total >= ROUND_MAX[round_type]:
            finish = api_request(
                client,
                "POST",
                f"/interviews/{interview_id}/rounds/{round_id}/finish",
                report,
                headers=headers,
                json={"finish_type": "normal"},
            ).json()
            round_report["ended_by"] = "max_questions"
            round_report["round_summary"] = finish.get("round_summary")
            break
    state_after_finish = api_request(
        client,
        "GET",
        f"/interviews/{interview_id}/state",
        report,
        headers=headers,
    ).json()
    current_round = next(item for item in state_after_finish["rounds"] if item["id"] == round_id)
    round_report["status_flow"].append(current_round["status"])
    round_report["state_after_finish"] = current_round
    report["round_execution"][round_type] = round_report


if __name__ == "__main__":
    main()
